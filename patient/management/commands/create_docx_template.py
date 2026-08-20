# patient/management/commands/create_docx_template.py
from pathlib import Path
from django.core.management.base import BaseCommand
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor


class Command(BaseCommand):
    help = "Génère le template DOCX avec mise en forme professionnelle (texte structuré)."

    def handle(self, *args, **options):
        app_dir = Path(__file__).resolve().parents[2]
        template_dir = app_dir / "docx_templates"
        template_dir.mkdir(parents=True, exist_ok=True)
        template_path = template_dir / "observation.docx"

        doc = Document()

        # 1. Style Normal
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.15

        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # 2. TITRE
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("{{ title }}")
        title_run.bold = True
        title_run.font.size = Pt(16)

        separator_para = doc.add_paragraph()
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_run = separator_para.add_run("_" * 80)
        separator_run.font.size = Pt(6)
        separator_run.font.color.rgb = RGBColor(100, 100, 100)

        # 3. BOUCLE SECTIONS
        doc.add_paragraph("{% for section in sections %}")

        section_title_para = doc.add_paragraph()
        section_title_para.paragraph_format.space_before = Pt(12)
        section_title_para.paragraph_format.space_after = Pt(4)
        section_title_run = section_title_para.add_run("{{ section.title }}")
        section_title_run.bold = True
        section_title_run.font.size = Pt(13)
        section_title_run.font.color.rgb = RGBColor(31, 73, 125)

        doc.add_paragraph("{% for block in section.blocks %}")

        block_title_para = doc.add_paragraph()
        block_title_para.paragraph_format.space_before = Pt(6)
        block_title_para.paragraph_format.space_after = Pt(2)
        block_title_run = block_title_para.add_run("{% if block.title %}{{ block.title }}{% endif %}")
        block_title_run.bold = True
        block_title_run.font.size = Pt(11)
        block_title_run.font.italic = True

        # 4. BOUCLE LIGNES (SIMPLE, SANS TABLEAU)
        doc.add_paragraph("{% for line in block.lines %}")

        line_para = doc.add_paragraph()
        line_para.paragraph_format.left_indent = Cm(0.5)
        line_para.paragraph_format.space_after = Pt(1)
        line_para.add_run("{{ line }}").font.size = Pt(10)

        doc.add_paragraph("{% endfor %}")

        doc.add_paragraph("{% endfor %}")  # fin for blocks

        section_sep_para = doc.add_paragraph()
        section_sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_sep_para.add_run("• • •").font.size = Pt(8)
        section_sep_para.add_run("• • •").font.color.rgb = RGBColor(150, 150, 150)

        doc.add_paragraph("{% endfor %}")  # fin for sections

        # 5. PIED DE PAGE
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.space_before = Pt(20)
        footer_run = footer_para.add_run("Document généré automatiquement par le système d'observation pédiatrique")
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.save(template_path)

        self.stdout.write(
            self.style.SUCCESS(
                f"Template DOCX généré avec succès : {template_path}"
            )
        )

        